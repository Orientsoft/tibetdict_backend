from fastapi import APIRouter, Body, Depends, HTTPException, UploadFile, File
from starlette.status import HTTP_400_BAD_REQUEST
from fastapi.responses import FileResponse
from io import BytesIO
from typing import List
from loguru import logger
from copy import deepcopy
import docx
import os, subprocess
import re
import platform
import shutil
import uuid
import json
from zipfile import ZipFile
from bs4 import BeautifulSoup

from model.user import User
from common.jwt import get_current_user_authorizer
from common.mongodb import AsyncIOMotorClient, get_database
from fastapi_plugins import depends_redis
from aioredis import Redis

from crud.file import create_file, get_file_list, count_file_by_query, delete_file, update_file, get_file, \
    create_upload_failed
from crud.work_history import count_work_history_by_query, get_work_history_id
from crud.word_dict import count_word_stat_dict_by_query, get_word_stat_dict_list
from model.file import FileCreateModel, OriginEnum, UploadFailedModel
from common.upload import MinioUploadPrivate
from common.common import contenttomd5, tokenize_sentence, judge_word
from common.search import bulk, delete_es_by_fileid
from config import ES_INDEX, timezone, SHARE_USER_ID
from datetime import datetime
from common.cache import get_cache, set_cache, del_cache
import sys
from common.search import query_es
from common.worker import celery_app
from model.work_history import WorkTypeEnum

router = APIRouter()
_platform = platform.system().lower()


class UploadError(Exception):
    def __init__(self):
        self.msg = '40017'


@router.get('/my/file', tags=['file'], name='我的文件')
async def get_my_file(user_id: str = None, search: str = None, is_check: bool = None,
                      user: User = Depends(get_current_user_authorizer()),
                      db: AsyncIOMotorClient = Depends(get_database)):
    query_obj = {'user_id': user_id}
    if search is not None:
        query_obj['file_name'] = {'$regex': search}
    if is_check is not None:
        query_obj['is_check'] = is_check
    data = await get_file_list(db, query_obj)
    count = await count_file_by_query(db, query_obj)
    return {
        'data': data,
        'count': count
    }


@router.get('/file/content', tags=['file'], name='某一文件内容')
async def get_file_content(file_id: str, is_origin: bool = False,
                           user: User = Depends(get_current_user_authorizer()),
                           db: AsyncIOMotorClient = Depends(get_database)):
    returnObj = {
        'content': [],
        'file_name': '',
        'is_check': None
    }
    db_file = await get_file(db, {'id': file_id})
    if not db_file:
        raise HTTPException(HTTP_400_BAD_REQUEST, '40011')
    # 既不是自己的文件，且不是共享文件
    if db_file.user_id != user.id and db_file.user_id != SHARE_USER_ID:
        raise HTTPException(HTTP_400_BAD_REQUEST, '40005')
    m = MinioUploadPrivate()
    content = m.get_object(db_file.parsed or db_file.origin if not is_origin else db_file.origin)
    temp_content = tokenize_sentence(content.decode('utf-8'))
    seq = 1
    for r in temp_content:
        if r.replace(' ', '') == '':
            continue
        returnObj['content'].append({'seq': seq, 'sentence': f"{r.strip()} "})
        seq = seq + 1
    returnObj['file_name'] = db_file.file_name
    returnObj['is_check'] = db_file.is_check
    return returnObj


@router.patch('/file', tags=['file'], name='修改文件')
async def patch_file(file_id: str = Body(...), content: str = Body(None), is_check: bool = Body(None),
                     book_name: str = Body(None),
                     author: str = Body(None),
                     version: str = Body(None),
                     tags: List = Body(None),
                     user: User = Depends(get_current_user_authorizer()),
                     db: AsyncIOMotorClient = Depends(get_database)):
    # 1.内容更新到minio，2.更新file的p_hash
    db_file = await get_file(db, {'id': file_id})
    if not db_file:
        raise HTTPException(HTTP_400_BAD_REQUEST, '40011')
    if db_file.user_id != user.id:
        raise HTTPException(HTTP_400_BAD_REQUEST, '40005')
    update_obj = {}
    if is_check is not None:
        if is_check and content is None and db_file.parsed is None:
            raise HTTPException(HTTP_400_BAD_REQUEST, '40018')
        update_obj['is_check'] = is_check
    if content is not None:
        m = MinioUploadPrivate()
        # 上传文件
        db_file.parsed = db_file.origin.replace('origin', 'parsed', 1)
        m.commit(content.encode('utf-8'), db_file.parsed)
        new_hash = contenttomd5(content.encode('utf-8'))
        update_obj['p_hash'] = new_hash
        update_obj['parsed'] = db_file.parsed
    if book_name:
        update_obj['book_name'] = book_name
    if author:
        update_obj['author'] = author
    if version:
        update_obj['version'] = version
    if tags:
        update_obj['tags'] = tags
    await update_file(db, {'id': file_id}, {'$set': update_obj})
    return {'msg': '2002'}


@router.delete('/file', tags=['file'], name='删除文件')
async def del_file(file_id: str,
                   user: User = Depends(get_current_user_authorizer()),
                   db: AsyncIOMotorClient = Depends(get_database), rd: Redis = Depends(depends_redis)):
    db_file = await get_file(db, {'id': file_id})
    if not db_file:
        raise HTTPException(HTTP_400_BAD_REQUEST, '40011')
    if db_file.user_id != user.id:
        raise HTTPException(HTTP_400_BAD_REQUEST, '40005')
    # 文件不能在work_history中出现
    use_count = await count_work_history_by_query(db, {'file_id': file_id})
    if use_count > 0:
        raise HTTPException(HTTP_400_BAD_REQUEST, '40012')
    m = MinioUploadPrivate()
    # m.remove(db_file.parsed)
    m.remove(db_file.origin)
    # 删除es中的数据
    delete_es_by_fileid(index=ES_INDEX, id=file_id)
    await delete_file(db, {'id': file_id, 'user_id': user.id})
    # 删除tree cache
    await del_cache(rd, user.id)
    return {'msg': '2002'}


@router.post('/file', tags=['file'], name='文件上传')
async def upload_file(file: UploadFile = File(...), path: str = Body(...), prefix_dir: str = Body(None),
                      user: User = Depends(get_current_user_authorizer()),
                      db: AsyncIOMotorClient = Depends(get_database), rd: Redis = Depends(depends_redis)):
    attr = file.filename.rsplit('.')[-1]
    if attr not in ['txt', 'docx', 'doc']:
        raise HTTPException(status_code=400, detail='40014')

    '''
    path:
    "doc/mm/asd.text"
    "doc/www.text"
    "doc/.DS_Store"
    "doc/__MACOSX/www.text"
    prefix_dir:"" "doc" ,"doc/sub" 
    '''

    # path 1.需删除第一层和末尾的文件名
    temp = path.split('/')
    filename = temp[-1]
    # 隐藏文件
    if filename.startswith('.'):
        raise HTTPException(status_code=HTTP_400_BAD_REQUEST, detail='40014')
    # MACOSX mac系统特有
    if 'MACOSX' in path:
        raise HTTPException(status_code=HTTP_400_BAD_REQUEST, detail='40014')
    subpath = '/'.join(temp[0:-1])
    # subpath case:1. mm/a   2. a  3. '' 4. .ipynb_checkpoints
    if '.' in subpath:
        raise HTTPException(status_code=HTTP_400_BAD_REQUEST, detail='40014')

    # file_name = '{}_{}'.format('_'.join(subpath.split('/')[2:]), file.filename) if subpath.split('/')[
    #                                                                                2:] else file.filename
    data = FileCreateModel(
        user_id=user.id,
        file_name=file.filename,
        is_check=False,
        tags=subpath.split('/')[0:2:] if subpath else []  # 前两级目录作为分类
    )

    if prefix_dir:
        if subpath:
            complete_path = f'{prefix_dir}/{subpath}/{data.id}.txt'
        else:
            complete_path = f'{prefix_dir}/{data.id}.txt'
    else:
        if subpath:
            complete_path = f'{subpath}/{data.id}.txt'
        else:
            complete_path = f'{data.id}.txt'

    # 原始文件
    data.origin = f'origin/{user.id}/{complete_path}'
    data.parsed = None
    # 查重
    result = await get_file(conn=db, query={'origin': data.origin})
    if result:
        raise HTTPException(status_code=HTTP_400_BAD_REQUEST, detail='40013')
    '''
    1.txt 本地存储，
    2.docx python-docx转换
    3.doc 不同平台不同方法，windows暂不支持
    '''
    try:
        origin_content = None
        if attr == 'txt':
            origin_content = file.file.read().decode('utf-8')
        elif attr in ['docx', 'doc']:
            # 临时目录 存储到本地
            if not os.path.exists('temp'):
                os.mkdir('temp')
            origin_temp_file_name = f"temp/{data.id}.{attr}"
            with open(origin_temp_file_name, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            if attr == 'docx':
                tmp = []
                try:
                    doc_file = docx.Document(origin_temp_file_name)
                    for para in doc_file.paragraphs:
                        tmp.append(para.text)
                    origin_content = '\n'.join(tmp)
                except:
                    tmp = []
                    document = ZipFile(origin_temp_file_name)
                    xml = document.read("word/document.xml")
                    wordObj = BeautifulSoup(xml.decode("utf-8"), features="lxml")
                    texts = wordObj.findAll("w:t")
                    for text in texts:
                        tmp.append(text.text)
                    origin_content = '\n'.join(tmp)
            elif attr == 'doc':
                saveas_txt_file_name = f"temp/{data.id}.txt"
                if _platform == 'linux':
                    cmd = f"catdoc {origin_temp_file_name} > {saveas_txt_file_name}"
                    os.system(cmd)
                elif _platform == 'darwin':
                    cmd = f"textutil -convert txt {origin_temp_file_name} -output {saveas_txt_file_name}"
                    os.system(cmd)
                else:
                    raise HTTPException(status_code=400, detail='40014')
                with open(saveas_txt_file_name, 'r', errors='ignore') as f:
                    origin_content = f.read()
                os.remove(saveas_txt_file_name)
            # 删除临时文件
            os.remove(origin_temp_file_name)

        m = MinioUploadPrivate()
        # es bulk操作
        actions = []
        temp_content = tokenize_sentence(origin_content)
        seq = 1
        for r in temp_content:
            if r.replace(' ', '') == '':
                continue
            actions.append({'index': {'_index': ES_INDEX, '_id': uuid.uuid1().hex}})
            actions.append({'id': data.id, 'seq': seq, 'content': r, 'user_id': user.id,
                            'createdAt': datetime.now(tz=timezone).isoformat()})
            seq = seq + 1
            if sys.getsizeof(actions) >= 25728:
                result = bulk(index=ES_INDEX, body=actions)
                if result['errors']:
                    logger.error(str(result))
                    raise UploadError
                else:
                    actions = []
        logger.info(sys.getsizeof(actions))
        result = bulk(index=ES_INDEX, body=actions)
        if result['errors']:
            logger.error(str(result))
            raise UploadError
    except Exception as e:
        logger.error(str(e))
        await create_upload_failed(conn=db, item=UploadFailedModel(
            path=path,
            user_id=user.id,
            file_id=data.id,
            prefix_dir=prefix_dir
        ))
        raise HTTPException(
            status_code=HTTP_400_BAD_REQUEST,
            detail='40017'
        )
    # 上传原始文件
    m.commit(origin_content.encode('utf-8'), data.origin)
    # parsed_content = re.sub(r"།(\s*)།", r"།།\r\n", origin_content)
    # 文件指纹
    data.o_hash = contenttomd5(origin_content.encode('utf-8'))
    await create_file(db, data)
    # 删除tree cache
    await del_cache(rd, user.id)
    return {'id': data.id}


@router.get('/tree', tags=['file'], name='文件目录')
async def get_my_content(origin: OriginEnum, user: User = Depends(get_current_user_authorizer()),
                         rd: Redis = Depends(depends_redis)):
    cache = await get_cache(rd, user.id if origin == OriginEnum.private else SHARE_USER_ID)
    if not cache:
        m = MinioUploadPrivate()
        if origin == OriginEnum.private:
            tree_data = m.list_tree(f'origin/{user.id}/')
        elif origin == OriginEnum.share:
            tree_data = m.list_tree(f'origin/{SHARE_USER_ID}/')
        else:
            raise HTTPException(HTTP_400_BAD_REQUEST, '40011')
        await set_cache(rd=rd, key=user.id if origin == OriginEnum.private else SHARE_USER_ID, info=tree_data)
    else:
        tree_data = cache
    return tree_data


@router.post('/content/file', tags=['file'], name='目录中内容')
async def get_content_file(origin: OriginEnum = Body(...), path: str = Body(None), search: str = Body(None),
                           user: User = Depends(get_current_user_authorizer()),
                           db: AsyncIOMotorClient = Depends(get_database)):
    m = MinioUploadPrivate()
    if origin == OriginEnum.private:
        user_id = user.id
    else:
        user_id = SHARE_USER_ID
    if path:
        comp_path = f'origin/{user_id}/{path}/'
    else:
        comp_path = f'origin/{user_id}/'
    result = m.list_content(comp_path, False)
    condition_file = []
    for item in result:
        condition_file.append(item['object_name'])
    query_obj = {'origin': {'$in': condition_file}}
    if search is not None:
        query_obj['file_name'] = {'$regex': search}
    data = await get_file_list(db, query_obj)
    count = await count_file_by_query(db, query_obj)
    return {
        'data': data,
        'count': count
    }


@router.post('/search', tags=['file'], name='搜索')
async def search_file(search: str = Body(...), origin: OriginEnum = Body(...), page: int = Body(1),
                      limit: int = Body(20),
                      user: User = Depends(get_current_user_authorizer())):
    start = (page - 1) * limit
    try:
        if search.endswith('་') or search.endswith('།'):
            search = search[:-1]

        queryObj = {
            "bool": {
                "must": [
                    # {"match_phrase": {"content": search}},
                    {"regexp": {"content": {"value": f".*{search}[འི|འུ|འོ|ས|ར]?[་| |།].*"}}},
                    {"term": {"user_id": user.id if origin == OriginEnum.private else SHARE_USER_ID}},
                ]
            }
        }
        result = query_es(index=ES_INDEX, queryObj=queryObj, start=start, size=limit)
    except Exception as e:
        logger.error(e)
        raise HTTPException(HTTP_400_BAD_REQUEST, '40017')
    returnObj = {
        'total': result['hits']['total']['value'],
        'data': []
    }
    # print(result)
    for r in result['hits']['hits']:
        content = []
        s = r['_source']['content'].replace(' ', '')
        search_start = s.find(search)
        search_end = search_start + len(search)
        header = s[0:s.find(search)]
        end = s[search_end:]
        content.append([header, search, end])
        returnObj['data'].append({
            'id': r['_source']['id'],
            'sentence': content,
            'seq': r['_source']['seq']
        })
    return returnObj


@router.post('/file/tokenize', tags=['file'], name='文件自动分词')
async def tokenize(file_ids: List[str] = Body(...), is_async: bool = Body(False),
                   user: User = Depends(get_current_user_authorizer()),
                   db: AsyncIOMotorClient = Depends(get_database)):
    data_file = await get_file_list(db, {'id': {'$in': file_ids}, 'is_check': False})
    for db_file in data_file:
        if db_file.user_id != user.id and user.id != SHARE_USER_ID:
            continue
        await update_file(db, {'id': db_file.id}, {'$set': {'tokenize_status': '0'}})
        resp = celery_app.send_task('worker:origin_tokenize', args=[db_file.id, user.id], queue='tibetan',
                                    routing_key='tibetan')
        if is_async is False and len(file_ids) == 1:
            resp = resp.get(timeout=300)
    return {'msg': '2001'}


@router.post('/file/search', tags=['file'], name='搜索某一文件内容')
async def search_file_content(file_id: str = Body(...), search: str = Body(...),
                              user: User = Depends(get_current_user_authorizer()),
                              db: AsyncIOMotorClient = Depends(get_database)):
    try:
        returnObj = {
            'content': [],
            'seq': [],
            'file_name': ''
        }
        db_file = await get_file(db, {'id': file_id})
        if not db_file:
            raise HTTPException(HTTP_400_BAD_REQUEST, '40011')
        # 既不是自己的文件，且不是共享文件
        if db_file.user_id != user.id and db_file.user_id != SHARE_USER_ID:
            raise HTTPException(HTTP_400_BAD_REQUEST, '40005')
        returnObj['file_name'] = db_file.file_name
        m = MinioUploadPrivate()
        content = m.get_object(db_file.origin)
        origin_content = content.decode('utf-8')
        temp_content = tokenize_sentence(origin_content)
        seq = 1
        for r in temp_content:
            if r.replace(' ', '') == '':
                continue
            returnObj['content'].append({'seq': seq, 'sentence': r})
            seq = seq + 1
        if search.endswith('་') or search.endswith('།'):
            search = search[:-1]
        queryObj = {
            "bool": {
                "must": [
                    # {"match_phrase": {"content": search}},
                    {"regexp": {"content": {"value": f".*{search}[འི|འུ|འོ|ས|ར]?[་| |།].*"}}},
                    {"term": {"id": file_id}},
                ]
            }
        }
        result = query_es(index=ES_INDEX, queryObj=queryObj, start=0, size=10000)
        for item in result['hits']['hits']:
            returnObj['seq'].append(item['_source']['seq'])
        return returnObj
    except Exception as e:
        logger.error(e)
        raise HTTPException(HTTP_400_BAD_REQUEST, '40017')


@router.post('/file/path/work_id', tags=['file'], name='某个目录中所有文件所属词频统计的的work_id')
async def get_work_id(origin: OriginEnum = Body(...), paths: List = Body(...), type: WorkTypeEnum = Body(...),
                      user: User = Depends(get_current_user_authorizer()),
                      db: AsyncIOMotorClient = Depends(get_database)):
    m = MinioUploadPrivate()
    if origin == OriginEnum.private:
        user_id = user.id
    else:
        user_id = SHARE_USER_ID
    condition = []
    for p in paths:
        comp_path = f'^origin/{user_id}{p}'
        regx = re.compile(comp_path, re.IGNORECASE)
        condition.append({'origin': {'$regex': regx}})
    work_id = await get_work_history_id(db, {'$or': condition, 'o_status': 1, 'work_type': type})
    return {
        'data': work_id
    }


@router.post('/file/tokenize/export', tags=['file'], name='辅助分词结果导出')
async def post_tokenize_export(ids: List[str] = Body(...), type: str = Body('new'),
                               user: User = Depends(get_current_user_authorizer(required=True)),
                               db: AsyncIOMotorClient = Depends(get_database)
                               ):
    m = MinioUploadPrivate()
    # 所有词
    words = []
    data_file = await get_file_list(db, {'id': {'$in': ids}, 'is_check': True})
    for db_file in data_file:
        if db_file.user_id != user.id and user.id != SHARE_USER_ID:
            continue
        try:
            content = m.get_object(db_file.parsed)
        except:
            continue
        temp_content = content.decode('utf-8').replace('\r', '').replace('\n', '').replace('\t', ' ').split(' ')
        words += temp_content
    # 藏语判断
    words = list(filter(judge_word, words))
    # 新词词库
    count = await count_word_stat_dict_by_query(db, {'type': 'used'})
    db_word_data = await get_word_stat_dict_list(db, {'type': 'used'}, 1, count)
    db_words = []
    for d in db_word_data:
        db_words.append(d.word)
    if type == 'new':
        export_word_list = list(set(words).difference(set(db_words)))
    else:
        export_word_list = list(set(words))
    if not os.path.exists('temp'):
        os.mkdir('temp')
    file_path = f'temp/tokenize-word-{datetime.now(tz=timezone).isoformat()[:10]}.txt'
    with open(file_path, 'w+', encoding='utf-8') as f:
        for w in export_word_list:
            f.writelines(f"{w}\n")
    headers = {'content-type': 'text/plain'}
    return FileResponse(file_path, headers=headers,
                        filename=f'tokenize-word-{datetime.now(tz=timezone).isoformat()[:10]}.txt')


@router.post('/file/once', tags=['file'], name='文件上传获取内容')
async def upload_file_get_content(file: UploadFile = File(...),
                                  user: User = Depends(get_current_user_authorizer()),
                                  db: AsyncIOMotorClient = Depends(get_database)):
    attr = file.filename.rsplit('.')[-1]
    if attr not in ['txt', 'docx', 'doc']:
        raise HTTPException(status_code=400, detail='40014')

    try:
        '''
        1.txt 本地存储，
        2.docx python-docx转换
        3.doc 不同平台不同方法，windows暂不支持
        '''
        temp_id = uuid.uuid1().hex
        origin_content = None
        if attr == 'txt':
            origin_content = file.file.read().decode('utf-8')
        elif attr in ['docx', 'doc']:
            # 临时目录 存储到本地
            if not os.path.exists('temp'):
                os.mkdir('temp')
            origin_temp_file_name = f"temp/{temp_id}.{attr}"
            with open(origin_temp_file_name, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            if attr == 'docx':
                tmp = []
                doc_file = docx.Document(origin_temp_file_name)
                for para in doc_file.paragraphs:
                    tmp.append(para.text)
                origin_content = '\n'.join(tmp)
            elif attr == 'doc':
                saveas_txt_file_name = f"temp/{temp_id}.txt"
                if _platform == 'linux':
                    cmd = f"catdoc {origin_temp_file_name} > {saveas_txt_file_name}"
                    os.system(cmd)
                elif _platform == 'darwin':
                    cmd = f"textutil -convert txt {origin_temp_file_name} -output {saveas_txt_file_name}"
                    os.system(cmd)
                else:
                    raise HTTPException(status_code=400, detail='40014')
                with open(saveas_txt_file_name, 'r') as f:
                    origin_content = f.read()
                os.remove(saveas_txt_file_name)
            # 删除临时文件
            os.remove(origin_temp_file_name)
        return {'data': origin_content}

    except Exception as e:
        logger.error(str(e))
        raise HTTPException(
            status_code=HTTP_400_BAD_REQUEST,
            detail='40014'
        )


@router.get('/work/tokenize', tags=['work'], name='分词记录')
async def get_my_tokenize_file(search: str = None,
                               user: User = Depends(get_current_user_authorizer()),
                               db: AsyncIOMotorClient = Depends(get_database)):
    query_obj = {'$or': [{'is_check': True}, {'tokenize_status': {'$in': ['0', '1', '2']}}], 'tokenize_user': user.id}
    if search is not None:
        query_obj['file_name'] = {'$regex': search}
    data = await get_file_list(db, query_obj)
    count = await count_file_by_query(db, query_obj)
    return {
        'data': data,
        'count': count
    }
